import pandas as pd
import os
import glob
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

# Define input and output folders
summary_folder = "summary"
report_folder = "reports-docs"
os.makedirs(report_folder, exist_ok=True)

# Function to generate DOCX reports
def generate_report(file_path):
    df = pd.read_csv(file_path)
    school_name = os.path.basename(file_path).replace("_summary.csv", "")
    total_learners = len(df)

    # Define literacy and numeracy categories
    literacy_categories = {
        "Below Expectations": ["Beginner (Lit)"],
        "Approaching Expectations": ["Letter", "Word"],
        "Meets Expectations": ["Paragraph"],
        "Above Expectations": ["Story", "Above (Lit)"]
    }
    numeracy_categories = {
        "Below Expectations": ["Beginner (Num)", "Count"],
        "Approaching Expectations": ["Number Rec.", "Addition"],
        "Meets Expectations": ["Subtraction", "Multiplication"],
        "Above Expectations": ["Division", "Above (Num)"]
    }

    # Function to categorize performance
    def categorize_performance(row, categories):
        for category, levels in categories.items():
            if row in levels:
                return category

    # Apply categorization
    df['Literacy Performance'] = df['Literacy Level'].apply(lambda x: categorize_performance(x, literacy_categories))
    df['Numeracy Performance'] = df['Numeracy Level'].apply(lambda x: categorize_performance(x, numeracy_categories))

    # Count learners per category
    literacy_counts = {category: df[df['Literacy Level'].isin(levels)].shape[0] for category, levels in literacy_categories.items()}
    numeracy_counts = {category: df[df['Numeracy Level'].isin(levels)].shape[0] for category, levels in numeracy_categories.items()}

    # Group by grade and performance
    table_literacy_counts = df.groupby(['Grade', 'Literacy Performance']).size().unstack(fill_value=0)
    table_numeracy_counts = df.groupby(['Grade', 'Numeracy Performance']).size().unstack(fill_value=0)

    # Calculate percentages
    literacy_percentages = {k: round((v / total_learners) * 100, 2) for k, v in literacy_counts.items()}
    numeracy_percentages = {k: round((v / total_learners) * 100, 2) for k, v in numeracy_counts.items()}

    # Create DOCX report
    doc = Document()
    doc.add_heading("SCHOOL BASELINE READING PERFORMANCE REPORT", level=1)
    doc.add_paragraph(f"School Name: {school_name}")
    doc.add_paragraph(f"Total Learners Assessed: {total_learners}")
    doc.add_paragraph("\nReading Performance Summary")

    # Literacy Table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "Number of Learners"
    hdr_cells[2].text = "Percentage of Total"

    for category in literacy_counts:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = str(literacy_counts[category])
        row_cells[2].text = f"{literacy_percentages[category]}%"

    doc.add_paragraph("\nKey Insights & Interpretation")
    for category, percentage in literacy_percentages.items():
        doc.add_paragraph(f"- {percentage}% of learners are in the {category} category.")

    doc.add_paragraph("\nRecommendations for Improvement")
    doc.add_paragraph("- Implement remedial reading sessions.")
    doc.add_paragraph("- Encourage home reading activities.")
    doc.add_paragraph("- Provide advanced reading materials for strong readers.")

    # Attach literature Graph Image
    graph_file_path = os.path.join(report_folder, f"{school_name}_literacy_graph.png")
    if os.path.exists(graph_file_path):
        doc.add_page_break()
        doc.add_paragraph("Literature Performance Summary Graph")
        doc.add_picture(graph_file_path, width=Inches(6))

    # Literacy performance table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Grade"
    hdr_cells[1].text = "Above expectations"
    hdr_cells[2].text = "Approaching Expectations"
    hdr_cells[3].text = "Below Expectations"
    hdr_cells[4].text = "Meets Expectations"

    for grade, row in table_literacy_counts.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(grade)
        row_cells[1].text = str(row.get('Above Expectations', 0))
        row_cells[2].text = str(row.get('Approaching Expectations', 0))
        row_cells[3].text = str(row.get('Below Expectations', 0))
        row_cells[4].text = str(row.get('Meets Expectations', 0))

    # Numeracy Section
    doc.add_page_break()
    doc.add_heading("SCHOOL BASELINE NUMERACY PERFORMANCE REPORT", level=1)
    doc.add_paragraph(f"School Name: {school_name}")
    doc.add_paragraph(f"Total Learners Assessed: {total_learners}")
    doc.add_paragraph("\nNumeracy Performance Summary")

    # Numeracy Table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "Number of Learners"
    hdr_cells[2].text = "Percentage of Total"

    for category in numeracy_counts:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = str(numeracy_counts[category])
        row_cells[2].text = f"{numeracy_percentages[category]}%"

    doc.add_paragraph("\nKey Insights & Interpretation")
    for category, percentage in numeracy_percentages.items():
        doc.add_paragraph(f"- {percentage}% of learners are in the {category} category.")

    doc.add_paragraph("\nRecommendations for Improvement")
    doc.add_paragraph("- Implement remedial numeracy sessions.")
    doc.add_paragraph("- Use hands-on activities to improve number sense.")
    doc.add_paragraph("- Provide advanced problem-solving exercises.")

    # Attach numeracy Graph Image
    graph_file_path = os.path.join(report_folder, f"{school_name}_numeracy_graph.png")
    if os.path.exists(graph_file_path):
        doc.add_page_break()
        doc.add_paragraph("Numeracy Performance Summary Graph")
        doc.add_picture(graph_file_path, width=Inches(6))

    # numeracy performance table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Grade"
    hdr_cells[1].text = "Above expectations"
    hdr_cells[2].text = "Approaching Expectations"
    hdr_cells[3].text = "Below Expectations"
    hdr_cells[4].text = "Meets Expectations"

    for grade, row in table_numeracy_counts.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(grade)
        row_cells[1].text = str(row.get('Above Expectations', 0))
        row_cells[2].text = str(row.get('Approaching Expectations', 0))
        row_cells[3].text = str(row.get('Below Expectations', 0))
        row_cells[4].text = str(row.get('Meets Expectations', 0))


    # Save DOCX report
    docx_file_path = os.path.join(report_folder, f"{school_name}_report.docx")
    doc.save(docx_file_path)
    print(f"✅ DOCX Report saved: {docx_file_path}")

# Process all summary files
summary_files = glob.glob(os.path.join(summary_folder, "*.csv"))
if not summary_files:
    print("❌ No summary files found in the 'summary' folder.")
else:
    for file in summary_files:
        generate_report(file)